import datetime
import os.path
import subprocess
from pathlib import Path
import shutil
from docx import Document
import re

current_dir = Path(__file__).resolve().parent
base_files_dir = os.path.join(current_dir, "base_files")
base_file_name = "base_letter.docx"
input_path = os.path.join(base_files_dir, base_file_name)
output_dir = os.path.join(current_dir, "out")
output_docx_name = "lettre_de_motivation.docx"
output_pdf_name = "lettre_de_motivation.pdf"

# Valeurs par défaut
default_replacements = {
    "salutations": "Monsieur, Madame,",
    "date": datetime.datetime.now().strftime("%d/%m/%y"),
    "continuer": "De plus, cela me permettrait une potentielle ouverture à un début de carrière au sein de ${nom_entreprise}.",
    "compétence_1": "développement logiciel",
    "compétence_2": "bases de données"
}


def detect_placeholders(text):
    """Detects placeholders in the format ${word} in the given text."""
    return [s for s in re.findall(r"\$\{(\w+)\}", text) if s.strip()]


def replace_placeholders(doc_path, replacements):
    """Replaces placeholders with user-provided values in the document while preserving style."""
    doc = Document(doc_path)

    # Iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Combine the text from all runs in the paragraph
        full_text = ''.join(run.text for run in paragraph.runs)

        # Replace the placeholders in the full text
        for placeholder, replacement in replacements.items():
            placeholder_pattern = f"${{{placeholder}}}"
            if placeholder_pattern in full_text:
                full_text = full_text.replace(placeholder_pattern, replacement)

        # Now update the runs with the new full_text
        runs = paragraph.runs
        start = 0
        for i, run in enumerate(runs):
            run_length = len(run.text)
            # Assign text without truncating and ensure the whole text is updated
            if i == len(runs) - 1:
                run.text = full_text[start:]
            else:
                run.text = full_text[start:start + run_length]
            start += run_length

    # Save the modified document after replacements
    doc.save(doc_path)
    print(f"[INFO] Placeholders replaced and file saved: {doc_path}")


def convert_docx_to_pdf(docx_path, pdf_path):
    """Converts the DOCX document to PDF using LibreOffice in headless mode."""
    try:
        # Run LibreOffice in headless mode (no GUI)
        subprocess.run([
            "C:\\Program Files\\LibreOffice\\program\\soffice.exe",
            "--headless",  # Runs without GUI
            "--convert-to", "pdf",
            "--outdir", os.path.dirname(pdf_path),
            docx_path
        ], check=True)
        print(f"[INFO] The file was successfully exported: {pdf_path}")
    except subprocess.CalledProcessError as e:
        print(f"[ERROR] The file could not be exported in pdf: {e}")


def create_letter():
    dir_name = input("[INPUT] Directory name: ")
    projet_dir = os.path.join(output_dir, dir_name)
    os.makedirs(projet_dir, exist_ok=True)
    print("[INFO] Directory absolute path:", projet_dir)

    docx_path = os.path.join(projet_dir, output_docx_name)
    pdf_path = os.path.join(projet_dir, output_pdf_name)

    shutil.copy(input_path, docx_path)

    doc = Document(docx_path)

    replacements = default_replacements.copy()

    whole_text = ""
    for paragraph in doc.paragraphs:
        whole_text += paragraph.text + '\n'

    placeholders = detect_placeholders(whole_text)

    visited_placeholders = set()

    while len(placeholders) > 0:
        for placeholder in placeholders:

            if placeholder in visited_placeholders:
                continue

            if placeholder in replacements:

                default_value = replacements[placeholder]

                user_input = input(
                    f"[INPUT] Replace ${placeholder} by (default value: '{default_value}', enter '$' to use it) : ")

                if user_input == "":
                    replacements[placeholder] = ""
                elif user_input == "$":
                    replacements[placeholder] = default_value
                elif user_input:
                    replacements[placeholder] = user_input

            else:
                replacements[placeholder] = input(
                    f"[INPUT] Replace ${placeholder} by: ")

            visited_placeholders.add(placeholder)

        replace_placeholders(docx_path, replacements)

        doc = Document(docx_path)

        whole_text = ""
        for paragraph in doc.paragraphs:
            whole_text += paragraph.text + '\n'

        placeholders = detect_placeholders(whole_text)

        convert_docx_to_pdf(docx_path, pdf_path)


if __name__ == "__main__":
    keep_going = ""
    while keep_going.upper() != "N":
        create_letter()
        keep_going = input("[INPUT] Continue? (enter 'n' to stop) : ")
